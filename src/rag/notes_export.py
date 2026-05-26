"""Export generated notes to PDF and DOCX."""

from __future__ import annotations

import io
import re
from pathlib import Path


def export_markdown_to_docx(markdown: str, *, title: str = "Notes") -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(title, 0)
    for line in (markdown or "").splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("•") or stripped.startswith("- "):
            doc.add_paragraph(stripped.lstrip("•- ").strip(), style="List Bullet")
        else:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.size = Pt(11)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_markdown_to_pdf(markdown: str, *, title: str = "Notes") -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.multi_cell(0, 8, _latin1_safe(title))
    pdf.ln(4)
    pdf.set_font("Helvetica", size=10)

    for line in (markdown or "").splitlines():
        stripped = line.strip()
        if not stripped:
            pdf.ln(3)
            continue
        if stripped.startswith("# "):
            pdf.set_font("Helvetica", "B", 12)
            pdf.multi_cell(0, 7, _latin1_safe(stripped[2:]))
            pdf.set_font("Helvetica", size=10)
        elif stripped.startswith("## "):
            pdf.set_font("Helvetica", "B", 11)
            pdf.multi_cell(0, 6, _latin1_safe(stripped[3:]))
            pdf.set_font("Helvetica", size=10)
        else:
            pdf.multi_cell(0, 5, _latin1_safe(stripped))
    out = pdf.output()
    return out if isinstance(out, bytes) else out.encode("latin-1")


def _latin1_safe(text: str) -> str:
    return text.encode("latin-1", errors="replace").decode("latin-1")
